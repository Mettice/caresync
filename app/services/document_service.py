import os
import uuid
import tempfile
from typing import Dict, List, Optional, Any
from fastapi import UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class DocumentService:
    """Service for processing and managing documents"""
    
    def __init__(self, vector_store_service):
        self.vector_store_service = vector_store_service
        self.document_store_path = os.getenv("DOCUMENT_STORE_PATH", "./data/documents")
        
        # Create document store directory if it doesn't exist
        os.makedirs(self.document_store_path, exist_ok=True)
    
    async def process_document(self, file: UploadFile, document_type: Optional[str] = None) -> Dict[str, Any]:
        """Process an uploaded document and add it to the vector store"""
        try:
            # Generate a unique document ID
            document_id = str(uuid.uuid4())
            
            # Get file extension
            filename = file.filename
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Save the uploaded file
            file_path = os.path.join(self.document_store_path, f"{document_id}{file_extension}")
            
            # Create a temporary file to store the uploaded content
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                # Read the uploaded file content
                content = await file.read()
                # Write content to the temporary file
                temp_file.write(content)
                temp_file.flush()
                
                # Extract text based on file type
                if file_extension == ".pdf":
                    documents = await self._process_pdf(temp_file.name, filename, document_id, document_type)
                elif file_extension == ".docx":
                    documents = await self._process_docx(temp_file.name, filename, document_id, document_type)
                else:
                    raise ValueError(f"Unsupported file type: {file_extension}")
                
                # Save the file to the document store
                with open(file_path, "wb") as dest_file:
                    # Reset file pointer to beginning
                    with open(temp_file.name, "rb") as src_file:
                        dest_file.write(src_file.read())
            
            # Clean up the temporary file
            os.unlink(temp_file.name)
            
            # Add documents to vector store
            num_chunks = await self.vector_store_service.add_documents(documents)
            
            return {
                "success": True,
                "document_id": document_id,
                "filename": filename,
                "document_type": document_type,
                "num_chunks": num_chunks,
                "metadata": {
                    "file_path": file_path,
                    "file_extension": file_extension
                }
            }
        
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    async def _process_pdf(self, file_path: str, filename: str, document_id: str, document_type: Optional[str]) -> List[Document]:
        """Process a PDF document and split it into chunks"""
        try:
            from PyPDF2 import PdfReader
            
            # Read PDF
            reader = PdfReader(file_path)
            text_content = ""
            
            # Extract text from each page
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\n\n=== Page {i+1} ===\n\n{page_text}"
            
            # Split text into chunks
            return self._split_text(text_content, filename, document_id, document_type)
        
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    async def _process_docx(self, file_path: str, filename: str, document_id: str, document_type: Optional[str]) -> List[Document]:
        """Process a DOCX document and split it into chunks"""
        try:
            import docx
            
            # Read DOCX
            doc = docx.Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            
            # Split text into chunks
            return self._split_text(text_content, filename, document_id, document_type)
        
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _split_text(self, text: str, filename: str, document_id: str, document_type: Optional[str]) -> List[Document]:
        """Split text into chunks for vector storage"""
        # Create text splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        # Split text into chunks
        chunks = text_splitter.split_text(text)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "filename": filename,
                    "document_id": document_id,
                    "chunk_id": i,
                    "document_type": document_type
                }
            )
            documents.append(doc)
        
        return documents